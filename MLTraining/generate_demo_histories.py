"""
Generates 5 synthetic patient profiles + one Word-document medical history each,
built from the same symptom-disease dataset used to train the model. Used as
demo/seed data for the capstone presentation (not real patient data).

Each document contains 2-3 dated "visit" entries drawn from real disease/symptom/
precaution/medication rows in the dataset, written as short clinical-style notes.
"""

import json
import random
from datetime import date, timedelta

import pandas as pd
from docx import Document

random.seed(7)

DATA_DIR = "data"
OUT_DIR = "demo_output"

DEMO_PATIENTS = [
    {"firstName": "Elena", "lastName": "Marchetti", "gender": "Female", "dob": "1988-03-14",
     "phone": "071234501", "email": "elena.marchetti@example.com", "address": "12 Birch Lane",
     "diseases": ["Fungal infection", "Migraine"]},
    {"firstName": "Marcus", "lastName": "Whitfield", "gender": "Male", "dob": "1975-11-02",
     "phone": "071234502", "email": "marcus.whitfield@example.com", "address": "45 Oak Street",
     "diseases": ["Hypertension", "Diabetes"]},
    {"firstName": "Priya", "lastName": "Nandakumar", "gender": "Female", "dob": "1993-07-21",
     "phone": "071234503", "email": "priya.nandakumar@example.com", "address": "8 Cedar Court",
     "diseases": ["Allergy", "Bronchial Asthma", "Common Cold"]},
    {"firstName": "Tomas", "lastName": "Novak", "gender": "Male", "dob": "1960-01-30",
     "phone": "071234504", "email": "tomas.novak@example.com", "address": "3 Willow Way",
     "diseases": ["Arthritis", "Osteoarthristis"]},
    {"firstName": "Amara", "lastName": "Okafor", "gender": "Female", "dob": "2001-09-09",
     "phone": "071234505", "email": "amara.okafor@example.com", "address": "27 Maple Ave",
     "diseases": ["Typhoid", "Gastroenteritis"]},
]


def load_reference_data():
    training = pd.read_csv(f"{DATA_DIR}/Training.csv")
    training = training.loc[:, ~training.columns.str.contains(r"^Unnamed")]
    symptom_columns = [c for c in training.columns if c != "prognosis"]

    import re
    def norm(name):
        return re.sub(r"\s+", " ", str(name).strip())

    training["prognosis"] = training["prognosis"].map(norm)

    description = pd.read_csv(f"{DATA_DIR}/description.csv")
    precautions = pd.read_csv(f"{DATA_DIR}/precautions_df.csv")
    medications = pd.read_csv(f"{DATA_DIR}/medications.csv")
    description["Disease"] = description["Disease"].map(norm)
    precautions["Disease"] = precautions["Disease"].map(norm)
    medications["Disease"] = medications["Disease"].map(norm)

    return training, symptom_columns, description, precautions, medications


def symptoms_for_disease(training, symptom_columns, disease):
    aliases = {"Peptic ulcer diseae": "Peptic ulcer disease"}
    lookup_name = disease
    for typo, correct in aliases.items():
        if correct == disease:
            lookup_name = typo
    row = training[training["prognosis"] == lookup_name].iloc[0]
    present = [s for s in symptom_columns if row[s] == 1]
    return [s.replace("_", " ").strip() for s in present]


def precautions_for_disease(precautions_df, disease):
    aliases = {"Peptic ulcer disease": "Peptic ulcer diseae"}
    lookup_name = aliases.get(disease, disease)
    match = precautions_df[precautions_df["Disease"] == lookup_name]
    if match.empty:
        return []
    row = match.iloc[0]
    cols = ["Precaution_1", "Precaution_2", "Precaution_3", "Precaution_4"]
    return [row[c] for c in cols if pd.notna(row[c]) and str(row[c]).strip()]


def medications_for_disease(medications_df, disease):
    match = medications_df[medications_df["Disease"] == disease]
    if match.empty:
        return []
    raw = match.iloc[0]["Medication"]
    try:
        return json.loads(str(raw).replace("'", '"'))
    except (json.JSONDecodeError, TypeError):
        return [raw] if pd.notna(raw) else []


def build_document(patient, training, symptom_columns, description, precautions_df, medications_df):
    doc = Document()
    doc.add_heading(f"Medical History — {patient['firstName']} {patient['lastName']}", level=1)
    doc.add_paragraph(f"Date of Birth: {patient['dob']}    Gender: {patient['gender']}")
    doc.add_paragraph("This document summarizes the patient's clinical visits to date.")

    visit_date = date.today() - timedelta(days=30 * len(patient["diseases"]) * 3)

    for disease in patient["diseases"]:
        symptoms = symptoms_for_disease(training, symptom_columns, disease)
        chosen_symptoms = random.sample(symptoms, k=min(4, len(symptoms)))
        desc_match = description[description["Disease"] == disease]
        desc_text = desc_match.iloc[0]["Description"] if not desc_match.empty else ""
        precs = precautions_for_disease(precautions_df, disease)
        meds = medications_for_disease(medications_df, disease)

        doc.add_heading(f"Visit — {visit_date.strftime('%B %Y')}", level=2)
        doc.add_paragraph(
            f"Chief complaint: {', '.join(chosen_symptoms)}."
        )
        doc.add_paragraph(f"Diagnosis: {disease}. {desc_text}")
        if precs:
            doc.add_paragraph("Precautions advised: " + ", ".join(precs) + ".")
        if meds:
            doc.add_paragraph("Medications prescribed: " + ", ".join(meds) + ".")

        visit_date = visit_date + timedelta(days=90)

    return doc


def main():
    import os
    os.makedirs(OUT_DIR, exist_ok=True)
    training, symptom_columns, description, precautions_df, medications_df = load_reference_data()

    manifest = []
    for patient in DEMO_PATIENTS:
        doc = build_document(patient, training, symptom_columns, description, precautions_df, medications_df)
        filename = f"{patient['firstName']}_{patient['lastName']}_history.docx".replace(" ", "_")
        path = f"{OUT_DIR}/{filename}"
        doc.save(path)
        manifest.append({**patient, "docxFile": filename})
        print(f"Generated {path} ({len(patient['diseases'])} visits: {', '.join(patient['diseases'])})")

    with open(f"{OUT_DIR}/manifest.json", "w") as f:
        json.dump(manifest, f, indent=2)
    print(f"\nManifest written to {OUT_DIR}/manifest.json")


if __name__ == "__main__":
    main()

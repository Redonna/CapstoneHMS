"""
Generates 5 synthetic demo patients with Word-doc medical histories built from
real symptom combinations sampled from the training dataset, so the AI symptom
summarizer has genuine overlapping vocabulary to work with when blending a
patient's past history into a new prediction.

Outputs:
  - demo_output/*.docx          (one history document per patient)
  - demo_output/seed_patients.sql  (ready-to-run T-SQL to insert everything)

This does NOT touch the database itself - run the generated .sql file with
sqlcmd separately so the insert step stays easy to inspect/re-run.
"""
import random
import uuid

import pandas as pd
from docx import Document

DATA_DIR = "data"
OUT_DIR = "demo_output"
UPLOADS_DIR = "../HospitalManagementSystem.API/Uploads/PatientHistory"

import os
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)

df = pd.read_csv(f"{DATA_DIR}/Diseases_and_Symptoms_dataset.csv")
symptom_cols = [c for c in df.columns if c != "diseases"]

PATIENTS = [
    {
        "first_name": "Anna", "last_name": "Bennett", "gender": "Female",
        "dob": "1990-04-12", "phone": "+1-555-0142", "email": "anna.bennett@example.com",
        "address": "12 Birchwood Lane, Springfield",
        "disease": "asthma",
    },
    {
        "first_name": "Maria", "last_name": "Torres", "gender": "Female",
        "dob": "1985-09-03", "phone": "+1-555-0198", "email": "maria.torres@example.com",
        "address": "45 Elm Street, Riverside",
        "disease": "acute sinusitis",
    },
    {
        "first_name": "David", "last_name": "Coleman", "gender": "Male",
        "dob": "1972-01-22", "phone": "+1-555-0176", "email": "david.coleman@example.com",
        "address": "8 Maple Avenue, Fairview",
        "disease": "gout",
    },
    {
        "first_name": "James", "last_name": "Carter", "gender": "Male",
        "dob": "1995-07-30", "phone": "+1-555-0164", "email": "james.carter@example.com",
        "address": "23 Cedar Court, Lakeview",
        "disease": "eczema",
    },
    {
        "first_name": "Sarah", "last_name": "Mitchell", "gender": "Female",
        "dob": "1988-11-15", "phone": "+1-555-0187", "email": "sarah.mitchell@example.com",
        "address": "17 Willow Drive, Brookside",
        "disease": "depression",
    },
]

VISIT_DATES = ["2026-05-04", "2026-06-18"]


def sample_symptoms(disease, n, rng):
    rows = df[df["diseases"] == disease]
    row = rows.sample(n=1, random_state=rng.randint(0, 1_000_000)).iloc[0]
    present = [c.replace("_", " ") for c in symptom_cols if row[c] == 1]
    rng.shuffle(present)
    return present[:n]


def build_document(patient, rng):
    doc = Document()
    doc.add_heading(f"Medical History - {patient['first_name']} {patient['last_name']}", level=1)

    visit1_symptoms = sample_symptoms(patient["disease"], 4, rng)
    doc.add_paragraph(f"Visit 1 - {VISIT_DATES[0]}")
    doc.add_paragraph(f"Chief complaint: {', '.join(visit1_symptoms)}.")
    doc.add_paragraph("Notes: Symptoms discussed, advised to monitor and return if they persist or worsen.")

    visit2_symptoms = sample_symptoms(patient["disease"], 6, rng)
    doc.add_paragraph(f"Visit 2 - {VISIT_DATES[1]}")
    doc.add_paragraph(f"Chief complaint: {', '.join(visit2_symptoms)}.")
    doc.add_paragraph(f"Diagnosis: {patient['disease'].title()}.")

    return doc


sql_lines = []

for patient in PATIENTS:
    rng = random.Random(hash(patient["email"]) & 0xFFFFFFFF)

    doc = build_document(patient, rng)
    file_stub = f"{patient['first_name']}_{patient['last_name']}_history"
    local_path = f"{OUT_DIR}/{file_stub}.docx"
    doc.save(local_path)

    stored_name = f"{uuid.uuid4()}.docx"
    doc.save(f"{UPLOADS_DIR}/{stored_name}")

    original_name = f"{file_stub}.docx"

    def esc(s):
        return s.replace("'", "''")

    sql_lines.append(f"""
GO
DECLARE @PatientId INT;
INSERT INTO Patients (FirstName, LastName, DateOfBirth, Gender, PhoneNumber, Email, Address, RegisteredAt, IsActive)
VALUES ('{esc(patient['first_name'])}', '{esc(patient['last_name'])}', '{patient['dob']}', '{patient['gender']}',
        '{patient['phone']}', '{patient['email']}', '{esc(patient['address'])}', GETUTCDATE(), 1);
SET @PatientId = SCOPE_IDENTITY();

INSERT INTO PatientHistoryEntries (PatientId, Title, Details, RecordDate, CreatedAt, AttachmentFileName, AttachmentStoredPath)
VALUES (@PatientId, 'Full Medical History', 'See attached document.', '{VISIT_DATES[1]}', GETUTCDATE(),
        '{esc(original_name)}', '{stored_name}');
""")

with open(f"{OUT_DIR}/seed_patients.sql", "w") as f:
    f.write("\n".join(sql_lines))

print(f"Generated {len(PATIENTS)} demo patients.")
print(f"Documents written to {OUT_DIR}/ and {UPLOADS_DIR}/")
print(f"SQL script written to {OUT_DIR}/seed_patients.sql")
